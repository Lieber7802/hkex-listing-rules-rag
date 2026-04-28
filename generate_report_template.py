from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

sections = doc.sections
for section in sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.5
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

title = doc.add_paragraph()
title_run = title.add_run('Phase 1 Implementation of Agentic RAG for HKEX Listing Rules Compliance')
title_run.bold = True
title_run.font.size = Pt(16)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('1st Interim Report')
subtitle_run.font.size = Pt(14)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

author = doc.add_paragraph()
author_run = author.add_run('Student Name: [Your Name]\nStudent ID: [Your Student ID]\nCourse: CS6520\nDate: [Date]')
author_run.font.size = Pt(12)
author.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

abstract_title = doc.add_paragraph()
abstract_run = abstract_title.add_run('Abstract')
abstract_run.bold = True
abstract_run.font.size = Pt(12)

abstract_para = doc.add_paragraph()
abstract_para.add_run('[This section should briefly summarize the project background, objectives, methodology, and preliminary results. Recommended length: 100-150 words. This section is optional for the 1st Interim Report but recommended for completeness.]')

doc.add_page_break()

toc_title = doc.add_paragraph()
toc_run = toc_title.add_run('Table of Contents')
toc_run.bold = True
toc_run.font.size = Pt(14)

toc_items = [
    '1. Introduction',
    '    1.1 Background and Problem Context',
    '    1.2 Project Objectives',
    '    1.3 Practical Value and Expected Outcome',
    '2. Related Work',
    '    2.1 Traditional Information Retrieval for Regulatory Documents',
    '    2.2 Neural and Dense Retrieval Methods',
    '    2.3 RAG and Agentic RAG Systems',
    '    2.4 Relationship Between Existing Work and This Project',
    '3. System Modeling and Structure',
    '    3.1 Problem Scope and System Boundary',
    '    3.2 Overall Architecture',
    '    3.3 LangGraph Workflow Structure',
    '    3.4 Design Justifications',
    '4. Methodology and Algorithms',
    '    4.1 Knowledge Base Construction Pipeline',
    '    4.2 Structure-Aware Chunking Algorithm',
    '    4.3 Hybrid Retrieval Strategy',
    '    4.4 Agentic Workflow: Planning, Retrieval and Reasoning',
    '    4.5 Current Limitations and Planned Enhancements',
    '5. Preliminary Performance Analysis or Experiments',
    '    5.1 Experimental Setup',
    '    5.2 Functional Verification',
    '    5.3 Preliminary Query Case Study',
    '    5.4 Preliminary Analysis',
    '6. Milestones and Overall Schedule',
    '    6.1 Work Completed So Far',
    '    6.2 Project Schedule and Milestones',
    '7. Work to be Completed for the Next Report',
    '    7.1 Technical Improvements',
    '    7.2 Evaluation and Reporting Work',
    '8. References'
]

for item in toc_items:
    p = doc.add_paragraph()
    p.add_run(item)

doc.add_page_break()

sec1 = doc.add_paragraph()
sec1_run = sec1.add_run('1. Introduction')
sec1_run.bold = True
sec1_run.font.size = Pt(14)

sec1_1 = doc.add_paragraph()
sec1_1_run = sec1_1.add_run('1.1 Background and Problem Context')
sec1_1_run.bold = True
sec1_1_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 250-300 words. Discuss the HKEX Listing Rules compliance context, the complexity of regulatory documents, the challenges users face in navigating these rules, and why traditional search methods are insufficient. Include statistics if available about the volume of rules, cross-references, and typical user queries.]')

sec1_2 = doc.add_paragraph()
sec1_2_run = sec1_2.add_run('1.2 Project Objectives')
sec1_2_run.bold = True
sec1_2_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 200-250 words. Clearly state the objectives of Phase 1: (1) Build a local backend prototype, (2) Implement document ingestion and structure-aware chunking, (3) Develop hybrid retrieval, (4) Create basic Agentic RAG workflow with planner and reasoning, (5) Generate citation-grounded answers. Also explicitly state what Phase 1 does NOT include: no web frontend, no Size Test Calculator tool, no benchmark dataset, no RAGAS evaluation, no production deployment.]')

sec1_3 = doc.add_paragraph()
sec1_3_run = sec1_3.add_run('1.3 Practical Value and Expected Outcome')
sec1_3_run.bold = True
sec1_3_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 200-250 words. Explain the practical value: (1) Assist compliance officers in quickly locating relevant rules, (2) Support multi-hop reasoning for complex compliance questions, (3) Provide traceable citations for audit purposes. Expected outcomes: (1) A working backend prototype, (2) A structured knowledge processing pipeline, (3) Preliminary experimental results, (4) Extensible interfaces for Phase 2 enhancements.]')

doc.add_page_break()

sec2 = doc.add_paragraph()
sec2_run = sec2.add_run('2. Related Work')
sec2_run.bold = True
sec2_run.font.size = Pt(14)

sec2_1 = doc.add_paragraph()
sec2_1_run = sec2_1.add_run('2.1 Traditional Information Retrieval for Regulatory Documents')
sec2_1_run.bold = True
sec2_1_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 250-300 words. Discuss traditional keyword-based search, BM25, and rule database search systems. Advantages: stable, fast, interpretable. Disadvantages: lack of semantic understanding, difficulty with multi-hop questions, inability to synthesize answers across multiple clauses. Cite 2-3 relevant works or systems.]')

sec2_2 = doc.add_paragraph()
sec2_2_run = sec2_2.add_run('2.2 Neural and Dense Retrieval Methods')
sec2_2_run.bold = True
sec2_2_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 250-300 words. Discuss dense embeddings, vector retrieval, and semantic matching approaches. Advantages: better semantic recall. Disadvantages: may retrieve semantically similar but legally imprecise content. Mention relevant models like BERT, BGE, and vector databases like FAISS, Chroma. Cite 2-3 relevant papers.]')

sec2_3 = doc.add_paragraph()
sec2_3_run = sec2_3.add_run('2.3 RAG and Agentic RAG Systems')
sec2_3_run.bold = True
sec2_3_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 300-350 words. Explain the difference between standard RAG (single-pass retrieval + generation) and Agentic RAG (planning, multi-round retrieval, evidence control). Discuss why Agentic RAG is more suitable for compliance Q&A. Mention relevant frameworks like LangChain, LangGraph, and recent Agentic RAG research. Cite 3-4 key papers or systems.]')

sec2_4 = doc.add_paragraph()
sec2_4_run = sec2_4.add_run('2.4 Relationship Between Existing Work and This Project')
sec2_4_run.bold = True
sec2_4_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 200-250 words. Explain how your system combines BM25 and dense retrieval, how it extends standard RAG with planner/router, conditional second retrieval, and citation-grounded answers. Also honestly acknowledge current limitations: MVP scope, simple planner, no tool/memory/evaluation yet. Position your work as a domain-specific Agentic RAG prototype for HKEX compliance.]')

doc.add_page_break()

sec3 = doc.add_paragraph()
sec3_run = sec3.add_run('3. System Modeling and Structure')
sec3_run.bold = True
sec3_run.font.size = Pt(14)

sec3_1 = doc.add_paragraph()
sec3_1_run = sec3_1.add_run('3.1 Problem Scope and System Boundary')
sec3_1_run.bold = True
sec3_1_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 200-250 words. Define the scope: focus on HKEX Listing Rules related to Notifiable Transactions, Connected Transactions, Size Tests, and disclosure/reporting obligations. Define input/output boundaries: natural language compliance questions as input, structured answers with citations as output. Mention what is explicitly out of scope for Phase 1.]')

sec3_2 = doc.add_paragraph()
sec3_2_run = sec3_2.add_run('3.2 Overall Architecture')
sec3_2_run.bold = True
sec3_2_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 250-300 words. Present the overall system architecture. Include a figure showing: Document Ingestion -> Cleaning -> Chunking -> Indexing -> Planner -> Retriever -> Reasoning -> Citation Formatter. Describe each major module briefly.]')

doc.add_paragraph('[Figure 1: System Architecture Diagram - Insert your architecture diagram here]')

sec3_3 = doc.add_paragraph()
sec3_3_run = sec3_3.add_run('3.3 LangGraph Workflow Structure')
sec3_3_run.bold = True
sec3_3_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 250-300 words. Describe the actual LangGraph StateGraph workflow: Planner -> Retriever -> Conditional Router -> Second Retrieval (optional) -> Reasoning -> Citation. Explain the role of each node and the conditional routing logic. Include a workflow diagram.]')

doc.add_paragraph('[Figure 2: LangGraph Workflow Diagram - Insert your workflow diagram here]')

sec3_4 = doc.add_paragraph()
sec3_4_run = sec3_4.add_run('3.4 Design Justifications')
sec3_4_run.bold = True
sec3_4_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 250-300 words. Justify key design choices: (1) Why structure-aware chunking instead of naive token-based splitting, (2) Why hybrid retrieval instead of dense-only, (3) Why LangGraph instead of hardcoded pipeline, (4) Why lightweight agentic approach instead of complex multi-agent system for Phase 1.]')

doc.add_page_break()

sec4 = doc.add_paragraph()
sec4_run = sec4.add_run('4. Methodology and Algorithms')
sec4_run.bold = True
sec4_run.font.size = Pt(14)

sec4_1 = doc.add_paragraph()
sec4_1_run = sec4_1.add_run('4.1 Knowledge Base Construction Pipeline')
sec4_1_run.bold = True
sec4_1_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 300-350 words. Describe the document ingestion pipeline: (1) Document loading from data/raw/, (2) Text cleaning and normalization, (3) Structure-aware chunking, (4) Metadata extraction (chapter, section, rule number), (5) Storage in data/processed/ and data/chunks/. Mention supported formats: .txt, .md, .pdf (basic support).]')

sec4_2 = doc.add_paragraph()
sec4_2_run = sec4_2.add_run('4.2 Structure-Aware Chunking Algorithm')
sec4_2_run.bold = True
sec4_2_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 300-350 words. Explain why naive token-based chunking is insufficient for legal documents. Describe the structure-aware approach: (1) Parse document hierarchy (chapter, section, rule), (2) Preserve rule numbers and cross-references, (3) Split long sections while maintaining context, (4) Store metadata: chunk_id, document_id, chapter, section_title, rule_number, text, source_path. Include a table showing chunk metadata fields.]')

doc.add_paragraph('[Table 1: Chunk Metadata Fields - Insert table here]')

sec4_3 = doc.add_paragraph()
sec4_3_run = sec4_3.add_run('4.3 Hybrid Retrieval Strategy')
sec4_3_run.bold = True
sec4_3_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 300-350 words. Describe the hybrid retrieval approach: (1) BM25 for lexical matching, (2) Dense retrieval using BGE-M3 embeddings via Ollama, (3) Score fusion strategy for combining results, (4) Why hybrid is more robust for legal text. Include equations or pseudo-code for score fusion if applicable.]')

sec4_4 = doc.add_paragraph()
sec4_4_run = sec4_4.add_run('4.4 Agentic Workflow: Planning, Retrieval and Reasoning')
sec4_4_run.bold = True
sec4_4_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 400-450 words. Detail the agentic workflow: (1) Planner: classify query as direct/multi-hop, generate sub-queries, decide on second retrieval, (2) Retriever: execute hybrid retrieval for query/sub-queries, (3) Conditional Router: determine if second retrieval is needed, (4) Reasoning Agent: synthesize answer from multiple evidence chunks, (5) Citation Formatter: generate structured citations. Include workflow pseudo-code or decision tree.]')

sec4_5 = doc.add_paragraph()
sec4_5_run = sec4_5.add_run('4.5 Current Limitations and Planned Enhancements')
sec4_5_run.bold = True
sec4_5_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 250-300 words. Honestly acknowledge limitations: (1) Planner is heuristic-based, not LLM-driven, (2) Second retrieval is not fully evidence-driven, (3) No tool integration yet, (4) No conversation memory, (5) No systematic evaluation. Outline planned enhancements for Phase 2: stronger planner, evidence coverage checking, tool interface, benchmark evaluation.]')

doc.add_page_break()

sec5 = doc.add_paragraph()
sec5_run = sec5.add_run('5. Preliminary Performance Analysis or Experiments')
sec5_run.bold = True
sec5_run.font.size = Pt(14)

sec5_1 = doc.add_paragraph()
sec5_1_run = sec5_1.add_run('5.1 Experimental Setup')
sec5_1_run.bold = True
sec5_1_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 200-250 words. Describe the experimental environment: (1) Hardware specifications, (2) Software stack: Python 3.10+, FastAPI, LangGraph, FAISS, BM25, (3) LLM: DeepSeek Reasoner, (4) Embeddings: BGE-M3 via Ollama, (5) Test document set: HKEX Listing Rules excerpts. Mention that this is a feasibility study, not a comprehensive benchmark.]')

sec5_2 = doc.add_paragraph()
sec5_2_run = sec5_2.add_run('5.2 Functional Verification')
sec5_2_run.bold = True
sec5_2_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 150-200 words. Report functional tests: (1) Document ingestion completes successfully, (2) Chunks are generated with correct metadata, (3) Indexes are built without errors, (4) API endpoints respond correctly, (5) Citations are traceable to source chunks. Mention test coverage: 40+ unit tests passing.]')

sec5_3 = doc.add_paragraph()
sec5_3_run = sec5_3.add_run('5.3 Preliminary Query Case Study')
sec5_3_run.bold = True
sec5_3_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 300-350 words. Present 2-4 example queries: (1) Direct clause retrieval: "What is Rule 14A.35?", (2) Multi-hop question: "What are the disclosure requirements for connected transactions and how do they differ from notifiable transactions?", (3) Calculation-related: "How do I calculate the size test threshold?". For each query, show: retrieved chunks, generated answer, citations. Use a table format for clarity.]')

doc.add_paragraph('[Table 2: Query Case Study Results - Insert table here]')

sec5_4 = doc.add_paragraph()
sec5_4_run = sec5_4.add_run('5.4 Preliminary Analysis')
sec5_4_run.bold = True
sec5_4_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 200-250 words. Analyze results: (1) System performs well on direct clause retrieval, (2) Multi-hop questions show reasonable synthesis but planner could be more sophisticated, (3) Citations are accurate and traceable, (4) Response time is acceptable for prototype. Identify areas for improvement: planner accuracy, evidence coverage checking, handling of ambiguous queries.]')

doc.add_page_break()

sec6 = doc.add_paragraph()
sec6_run = sec6.add_run('6. Milestones and Overall Schedule')
sec6_run.bold = True
sec6_run.font.size = Pt(14)

sec6_1 = doc.add_paragraph()
sec6_1_run = sec6_1.add_run('6.1 Work Completed So Far')
sec6_1_run.bold = True
sec6_1_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 200-250 words. List completed work: (1) Project skeleton and configuration system, (2) Document ingestion pipeline (loader, cleaner, chunker), (3) Hybrid retrieval (BM25, FAISS, embedder), (4) LangGraph workflow with planner and reasoning agents, (5) Citation formatter, (6) FastAPI interface with /health and /chat endpoints, (7) Unit tests (40+ tests), (8) Documentation (README, README-zh, spec.md).]')

sec6_2 = doc.add_paragraph()
sec6_2_run = sec6_2.add_run('6.2 Project Schedule and Milestones')
sec6_2_run.bold = True
sec6_2_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 150-200 words. Present a timeline or table showing: (1) Phase 1 completion (current), (2) Phase 2 planning, (3) Tool integration, (4) Evaluation and benchmark, (5) Final report preparation. Use a Gantt chart or milestone table.]')

doc.add_paragraph('[Table 3: Project Milestones - Insert table here]')

doc.add_page_break()

sec7 = doc.add_paragraph()
sec7_run = sec7.add_run('7. Work to be Completed for the Next Report')
sec7_run.bold = True
sec7_run.font.size = Pt(14)

sec7_1 = doc.add_paragraph()
sec7_1_run = sec7_1.add_run('7.1 Technical Improvements')
sec7_1_run.bold = True
sec7_1_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 200-250 words. Outline planned technical improvements: (1) Enhance planner with LLM-based intent classification and task decomposition, (2) Implement evidence coverage checker, (3) Add answer verification mechanism, (4) Integrate tool interface (Size Test Calculator, Rule Lookup Tool), (5) Improve retrieval with query rewriting and reranking.]')

sec7_2 = doc.add_paragraph()
sec7_2_run = sec7_2.add_run('7.2 Evaluation and Reporting Work')
sec7_2_run.bold = True
sec7_2_run.font.size = Pt(12)

doc.add_paragraph('[Recommended: 150-200 words. Describe evaluation plans: (1) Build benchmark dataset with 20-30 annotated questions, (2) Implement evaluation metrics (retrieval recall, citation quality, answer correctness), (3) Compare with baseline systems, (4) Prepare final system demonstration, (5) Complete final report with comprehensive results and analysis.]')

doc.add_page_break()

sec8 = doc.add_paragraph()
sec8_run = sec8.add_run('8. References')
sec8_run.bold = True
sec8_run.font.size = Pt(14)

doc.add_paragraph('[Include 12-20 references in IEEE format. Suggested categories:]')

doc.add_paragraph('[1-3] RAG and Agentic RAG papers')
doc.add_paragraph('[4-6] Dense retrieval and embedding models')
doc.add_paragraph('[7-9] BM25 and hybrid retrieval methods')
doc.add_paragraph('[10-12] LangChain and LangGraph documentation')
doc.add_paragraph('[13-15] HKEX official documentation and rules')
doc.add_paragraph('[16-18] DeepSeek, BGE-M3, FAISS technical references')
doc.add_paragraph('[19-20] Related compliance or legal QA systems')

doc.add_paragraph()
doc.add_paragraph('[Example format:]')
doc.add_paragraph('[1] Y. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks," in Advances in Neural Information Processing Systems, 2020.')

doc.save('D:\\CItyUFile\\CS6520 Project\\1st_Interim_Report_Template.docx')
print('Report template generated successfully!')
